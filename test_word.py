from docx import Document

document = Document()

document.add_heading('わーどさんぷる', 0)
document.add_paragraph('てすとさんぷる')

document.save('sample.docx')