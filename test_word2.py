from docx import Document
from docx.shared import Inches

# 1. 「sample.docx」を読み込んで画像を貼り付ける
doc = Document('sample.docx')
doc.add_picture('image.png', width=Inches(2), height=Inches(2))

# 2. 「sample.docx」内の文字数をカウントして出力する
text_count = 0
for paragraph in doc.paragraphs:
    text_count += len(paragraph.text)
print("文字数:", text_count)

# 3. 「sample_answer.docx」としてドキュメントを保存する
doc.save('sample_answer.docx')